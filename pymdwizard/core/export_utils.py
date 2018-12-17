#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''
The MetadataWizard (pymdwizard) software was developed by the
U.S. Geological Survey Fort Collins Science Center.
See: https://github.com/usgs/fort-pymdwizard for current project source code
See: https://usgs.github.io/fort-pymdwizard/ for current user documentation
See: https://github.com/usgs/fort-pymdwizard/tree/master/examples
for examples of use in other scripts

Authors: Colin Talbert, USGS and Michelle M. Fink, CNHP
    with contributed code from:
https://stackoverflow.com/questions/45819839/when-using-python-docx-how-to-enable-spelling-in-output-document
https://github.com/python-openxml/python-docx/issues/74

License:   Creative Commons Attribution 4.0 International (CC BY 4.0)
           http://creativecommons.org/licenses/by/4.0/

PURPOSE
------------------------------------------------------------------------------
Convert FGDC XML documents to a rendered Word format (docx)
based on
https://github.com/talbertc-usgs/Notebooks/Metadata/Convert FGDC XMLs to Formatted Word docx.ipynb

SCRIPT DEPENDENCIES
------------------------------------------------------------------------------
    This script is part of a modified pymdwizard package and is not intended to
    be used independently.  All pymdwizard package requirements are needed.

    See imports section for external packages used in this script as well as
    inter-package dependencies

DISCLAIMER
------------------------------------------------------------------------------
USGS:
This code has NOT been reviewed or endorsed by the U.S. Geological Survey
(USGS), but was incorporated by a third party (Michelle M. Fink, CNHP). No
warranty, expressed or implied, is made by the USGS or the U.S. Government as
to the functionality of the software and related material nor shall the fact
of release constitute any such warranty. Furthermore, the software is released
on condition that neither the USGS nor the U.S. Government shall be held
liable for any damages resulting from its authorized or unauthorized use.

CNHP:
See Section 5 of the license for Disclaimer of Warranties and Limitation of
Liability. This disclaimer applies to the authors, The Colorado Natural Heritage
Program (CNHP), Colorado State University, and the State of Colorado.
'''

#TODO: Right now this just exports to a Word docx, would like to include other formats

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from pymdwizard.core import review_utils
from pymdwizard.core.review_utils import _get_longname, _add_child_content
from pymdwizard.core import xml_utils

def add_bookmark(paragraph, bookmark_name):
    ''' Adds a Microsoft Word bookmark to a paragraph '''
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.append(start)

    text = OxmlElement('w:r')
    tag.append(text)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    end.set(qn('w:name'), bookmark_name)
    tag.append(end)

    return run

def add_hyperlink(paragraph, link_to, text, is_external):
    ''' Adds a hyperlink within a paragraph to an internal bookmark
    or an external url '''

    part = paragraph.part

    hyperlink = OxmlElement('w:hyperlink')
    if is_external:
        r_id = part.relate_to(link_to, RT.HYPERLINK, is_external=is_external)
        hyperlink.set(qn('r:id'), r_id, )
    else:
        hyperlink.set(qn('w:anchor'), link_to, )

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def _add_abbrev_child_content(doc, node, indent=0.25):
    ''' Adds all the content of children of 'node' without
    adding 'node' itself '''
    for child in node.children:
        _add_child_content(doc, child, indent+0.25)

def md_to_docx(xml_fname, docx_fname):
    ''' Converts FGDC xml metadata to Microsoft Word docx '''
    document = Document()
    #Set narrow margins:
    sections = document.sections
    section = sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    DOCX = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    #This prevents Word from ignoring spelling errors:
    element = document.settings.element.find(DOCX + 'proofState')
    element.attrib[DOCX + 'grammar'] = 'dirty'
    element.attrib[DOCX + 'spelling'] = 'dirty'

    #My tweaked custom styles, in my tweaked fort-pymwizard (mmf branch)
    review_utils._load_mmf_styles(document)

    md = xml_utils.XMLRecord(xml_fname)
    #Set the citation title as the document title
    md_title = md.metadata.idinfo.citation.citeinfo.title.text
    title1 = document.add_heading(md_title, level=1)
    title1.style = document.styles['fgdc title']
    title2 = document.add_heading('Metadata:', level=2)
    title2.style = document.styles['fgdc heading 2']

    #Populate a 'table of contents' (of sorts) at beginning of doc
    for child in md.metadata.children:
        long_name = _get_longname(child.tag)
        #If there is no long name, skip
        if long_name != child.tag:
            link = document.add_paragraph(style='fgdc link')
            add_hyperlink(link, long_name.replace(' ', '_'), long_name, is_external=False)
            link.paragraph_format.left_indent = Inches(0.25)
            link.paragraph_format.line_spacing = 1

    #Set a heading for each top child, then recursively populate with all children
    for child in md.metadata.children:
        long_name = _get_longname(child.tag)
        #Don't make a heading if there is no long name
        if long_name != child.tag:
            #Separate sections with a vertical line:
            document.add_paragraph('_'*72, style='fgdc bar')
            #Add section heading and link to toc:
            section_title = document.add_heading(long_name+ ':', level=3)
            add_bookmark(section_title, long_name.replace(' ', '_'))
            section_title.style = document.styles['fgdc heading 3']
            section_title.paragraph_format.space_after = Inches(.15)
        #Add remaining tags and content
        _add_abbrev_child_content(document, child, indent=0.1)

    document.save(docx_fname)
